import os
import cv2
import easyocr
import re
import unicodedata
import numpy as np
import tkinter as tk
from tkinter import messagebox
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg
from PIL import Image
import fitz  # PyMuPDF
from docx2pdf import convert
from tkinter import filedialog
import docx

# ==========================================
# 1. CẤU HÌNH HỆ THỐNG & AI GLOBAL (LAZY LOAD)
# ==========================================
AI_MODELS = {'easyocr': None, 'vietocr': None}


def get_ai_models():
    """Tải mô hình AI một lần duy nhất."""
    if AI_MODELS['easyocr'] is None:
        print("\n--- ĐANG NẠP MÔ HÌNH AI (EASYOCR + VIETOCR)... ---")
        AI_MODELS['easyocr'] = easyocr.Reader(['vi', 'en'], gpu=True)

        config = Cfg.load_config_from_name('vgg_transformer')
        config['device'] = 'cpu'
        AI_MODELS['vietocr'] = Predictor(config)
        print("--- NẠP MÔ HÌNH THÀNH CÔNG! ---")

    return AI_MODELS['easyocr'], AI_MODELS['vietocr']


# ==========================================
# 2. HÀM XỬ LÝ CHÍNH
# ==========================================

def process_and_redact(image_path, output_path, parent_window):
    """
    Nhận ảnh gốc, OCR, mở menu chọn vùng cần che, xử lý và lưu ảnh.
    Trả về True nếu thành công, False nếu người dùng hủy hoặc lỗi.
    """
    reader, vietocr_reader = get_ai_models()

    # --- 1. ĐỌC ẢNH HỖ TRỢ ĐƯỜNG DẪN TIẾNG VIỆT ---
    try:
        img_array = np.fromfile(image_path, np.uint8)
        img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

        if img is None:
            pil_img = Image.open(image_path).convert('RGB')
            img = np.array(pil_img)
            img = img[:, :, ::-1].copy()
    except Exception as e:
        messagebox.showerror(
            "Lỗi định dạng",
            f"Không thể đọc file này như một bức ảnh.\nChi tiết: {str(e)}",
            parent=parent_window,
        )
        return False

    if img is None:
        messagebox.showerror(
            "Lỗi định dạng",
            "Hệ thống không nhận diện được file này là ảnh.",
            parent=parent_window,
        )
        return False

    img_h, img_w = img.shape[:2]

    # --- 2. TIỀN XỬ LÝ ẢNH DÀNH CHO OCR ---
    gray_img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray_img.shape
    max_size = 1200
    scale_factor = max_size / max(h, w) if max(h, w) > max_size else 1.5

    gray_img = cv2.resize(gray_img, None, fx=scale_factor, fy=scale_factor, interpolation=cv2.INTER_CUBIC)
    blurred_img = cv2.GaussianBlur(gray_img, (3, 3), 0)

    print("🔍 Hệ thống đang quét tọa độ vùng chữ...")
    pre_results = reader.readtext(blurred_img)

    print("\n===== EASY OCR ĐỌC ĐƯỢC =====")
    for idx, (bbox, text, conf) in enumerate(pre_results, start=1):
        print(f"{idx}. Text: {text} | Conf: {conf:.2f} | BBox: {bbox}")
    print("===== HẾT EASY OCR =====\n")

    print("🔍 Hệ thống đang phân tích loại giấy tờ để đưa ra đề xuất...")
    all_ocr_text = " ".join([res[1].lower() for res in pre_results])
    all_ocr_text = unicodedata.normalize('NFD', all_ocr_text)
    all_ocr_text = "".join(ch for ch in all_ocr_text if unicodedata.category(ch) != 'Mn')

    # ==========================================
    # HÀM HỖ TRỢ CƠ BẢN
    # ==========================================
    def remove_accents(text):
        return ''.join(ch for ch in unicodedata.normalize('NFD', text) if unicodedata.category(ch) != 'Mn')

    def normalize_text(text):
        return re.sub(r'\s+', ' ', remove_accents(text).lower()).strip()

    def has_keyword(norm_text, keywords):
        return any(re.search(r'\b' + re.escape(kw) + r'\b', norm_text) for kw in keywords)

    def redact_rect(img_ref, x1, y1, x2, y2, padding=2):
        img_h_ref, img_w_ref = img_ref.shape[:2]
        cv2.rectangle(
            img_ref,
            (max(0, int(x1) - padding), max(0, int(y1) - padding)),
            (min(img_w_ref, int(x2) + padding), min(img_h_ref, int(y2) + padding)),
            (0, 0, 0),
            -1,
        )

    def pixelate_region(img_ref, x1, y1, x2, y2, blocks=12):
        x1, y1 = max(0, int(x1)), max(0, int(y1))
        x2, y2 = min(img_ref.shape[1], int(x2)), min(img_ref.shape[0], int(y2))
        roi = img_ref[y1:y2, x1:x2]
        if roi.size == 0:
            return
        h_roi, w_roi = roi.shape[:2]
        temp = cv2.resize(roi, (blocks, blocks), interpolation=cv2.INTER_LINEAR)
        img_ref[y1:y2, x1:x2] = cv2.resize(temp, (w_roi, h_roi), interpolation=cv2.INTER_NEAREST)

    def looks_like_license_plate_text(text):
        text = unicodedata.normalize('NFD', text.upper())
        text = ''.join(ch for ch in text if unicodedata.category(ch) != 'Mn')
        compact = re.sub(r'[^A-Z0-9]', '', text)

        return (
            re.search(r'\d{2}[A-Z]{1,2}\d?', compact)
            and len(re.findall(r'\d', compact)) >= 5
            and 6 <= len(compact) <= 12
        )

    ocr_texts = [res[1] for res in pre_results]
    has_plate_candidate = any(looks_like_license_plate_text(t) for t in ocr_texts)

    for i in range(len(ocr_texts) - 1):
        if looks_like_license_plate_text(ocr_texts[i] + ocr_texts[i + 1]):
            has_plate_candidate = True
            break

    # --- 3. LOGIC AI ĐỀ XUẤT TỰ ĐỘNG ---
    suggestions = {
        'face': False,
        'barcode': False,
        'finger': False,
        'qr': False,
        'id_num': False,
        'name': False,
        'dob': False,
        'address': False,
        'cv_contact': False,
        'plate': False,
        'driver_license_num': False,
    }

    document_type = "Không xác định (Tự chọn)"

    if 'can cuoc' in all_ocr_text or 'citizen' in all_ocr_text:
        document_type = "Căn cước công dân (Mặt trước)"
        suggestions.update({
            'face': True,
            'qr': True,
            'id_num': True,
            'name': True,
            'dob': True,
            'address': True,
        })
    elif 'dac diem nhan dang' in all_ocr_text or 'personal identification' in all_ocr_text:
        document_type = "Căn cước công dân (Mặt sau)"
        suggestions.update({'finger': True, 'id_num': True, 'name': True})
    elif 'bao hiem y te' in all_ocr_text or 'bhyt' in all_ocr_text or 'the bhyt' in all_ocr_text:
        document_type = "Thẻ Bảo hiểm y tế"
        suggestions.update({
            'id_num': True,
            'name': False,
            'dob': False,
            'barcode': False,
        })
    elif 'sinh vien' in all_ocr_text or 'student' in all_ocr_text:
        document_type = "Thẻ sinh viên / Thẻ học sinh"
        suggestions.update({'id_num': True, 'name': True, 'barcode': True, 'face': True, 'dob': True})
    elif any(k in all_ocr_text for k in ['ho so', 'cv', 'resume', 'lien he', 'email', 'dien thoai', 'hoc van', 'kinh nghiem', 'ky nang']):
        document_type = "CV / Hồ sơ cá nhân"
        suggestions.update({'face': True, 'cv_contact': True})
    elif any(k in all_ocr_text for k in ['giay phep lai xe', 'gplx', 'so giay phep', 'ngay cap', 'noi cap']):
        document_type = "Giấy phép lái xe"
        suggestions.update({'driver_license_num': True})
    elif any(k in all_ocr_text for k in ['dang ky xe', 'chung nhan dang ky xe']):
        document_type = "Đăng ký xe"
        suggestions.update({'plate': True})
    elif has_plate_candidate:
        document_type = "Biển số xe"
        suggestions.update({'plate': True})

    if document_type == "Không xác định (Tự chọn)":
        document_type = "Ảnh thường / Ảnh chung"
        suggestions.update({'face': True, 'cv_contact': True, 'id_num': True, 'name': True, 'dob': True,})

    print(f"📌 HỆ THỐNG ĐỀ XUẤT: [{document_type}]")

    # --- 4. GIAO DIỆN XÁC NHẬN ---
    user_choices = {}
    is_submitted = [False]

    dialog = tk.Toplevel(parent_window)
    dialog.title("Hệ thống che thông tin ")
    dialog.geometry("460x500")
    dialog.configure(bg="#F0F7F4")
    dialog.transient(parent_window)
    dialog.grab_set()

    choice_vars = {
        key: tk.BooleanVar(value=value)
        for key, value in suggestions.items()
    }

    tk.Label(
        dialog,
        text=f"🔍 Hệ thống nhận diện: {document_type}",
        bg="#E8F1F5",
        fg="#1D3557",
        font=("Arial", 11, "italic"),
        bd=1,
        relief="solid",
        padx=10,
        pady=5,
    ).pack(fill="x", padx=30, pady=(15, 5))

    tk.Label(
        dialog,
        text="🛠️ ĐIỀU CHỈNH VÙNG CẦN CHE:",
        bg="#F0F7F4",
        fg="#2C3E50",
        font=("Arial", 11, "bold"),
    ).pack(pady=5)

    f1 = tk.LabelFrame(dialog, text=" Vùng hình ảnh & Mã vạch ", bg="#F0F7F4", font=("Arial", 10, "bold"), padx=10, pady=5)
    f1.pack(fill="x", padx=30, pady=5)

    f2 = tk.LabelFrame(dialog, text=" Thông tin văn bản ", bg="#F0F7F4", font=("Arial", 10, "bold"), padx=10, pady=5)
    f2.pack(fill="x", padx=30, pady=5)

    image_options = [
        ('face', "Che khuôn mặt"),
        ('barcode', "Che mã vạch"),
        ('finger', "Che dấu vân tay"),
        ('qr', "Che mã QR"),
        ('plate', "Che biển số xe"),
    ]

    text_options = [
        ('id_num', "Che Số định danh/Số thẻ"),
        ('name', "Che Họ và tên"),
        ('dob', "Che Ngày sinh"),
        ('address', "Che Địa chỉ"),
        ('cv_contact', "Che Email / Số điện thoại tự do"),
        ('driver_license_num', "Che số giấy phép lái xe"),
    ]

    def make_redaction_checkbox(parent, key, label):
        is_allowed = suggestions.get(key, False)

        cb = tk.Checkbutton(
            parent,
            text=label,
            variable=choice_vars[key],
            bg="#F0F7F4",
            activebackground="#F0F7F4",
            state="normal" if is_allowed else "disabled",
        )

        cb.pack(anchor="w", padx=20)
        return cb

    for key, label in image_options:
        make_redaction_checkbox(f1, key, label)

    for key, label in text_options:
        make_redaction_checkbox(f2, key, label)

    def on_submit():
        for key, var in choice_vars.items():
            user_choices[key] = var.get()
        is_submitted[0] = True
        dialog.destroy()

    tk.Button(
        dialog,
        text="Xác nhận và Xử lý ảnh",
        command=on_submit,
        bg="#A8DADC",
        fg="#1D3557",
        font=("Arial", 11, "bold"),
        padx=25,
        pady=5,
        cursor="hand2",
    ).pack(pady=15)

    parent_window.wait_window(dialog)

    if not is_submitted[0]:
        print("Đã hủy thao tác xử lý ảnh.")
        return False

    # --- 5. SETUP TỪ KHÓA TÌM KIẾM ---
    sensitive_keywords = []
    next_line_keywords = []
    vietnam_address_words = []

    if user_choices['id_num']:
        sensitive_keywords.extend([
            'so', 'so the', 'so cccd', 'so ho chieu', 'ma so thue', 'mst',
            'so bhxh', 'so bhyt', 'ma so', 'ma the', 'no'
        ])

    if user_choices['dob']:
        sensitive_keywords.extend([
            'ngay sinh', 'sinh ngay', 'date of birth', 'dob', 'nam sinh',
            'ngay sinh:', 'ngay/thang/nam sinh'
        ])

    if user_choices['address']:
        sensitive_keywords.extend([
            'que quan', 'place of origin', 'noi sinh', 'place of birth',
            'noi thuong tru', 'thuong tru', 'place of residence', 'residence',
            'dia chi', 'address', 'noi o hien tai', 'noi dki kcb ban dau', 'noi dk kcb'
        ])
        vietnam_address_words.extend([
            'tp hcm', 'tphcm', 'ho chi minh', 'ha noi', 'da nang', 'can tho',
            'hai phong', 'binh duong', 'dong nai', 'long an', 'ba ria', 'vung tau',
            'quan', 'huyen', 'thi xa', 'thanh pho', 'phuong', 'xa', 'thi tran',
            'ap', 'thon', 'khu pho', 'duong', 'so nha'
        ])

    if user_choices['cv_contact']:
        sensitive_keywords.extend(['dien thoai', 'so dien thoai', 'phone', 'mobile', 'tel'])
        next_line_keywords.extend(['dien thoai', 'so dien thoai', 'phone', 'mobile', 'tel'])

    # ==========================================
    # HÀM HỖ TRỢ REDACTION
    # ==========================================
    def line_has_sensitive_pattern(text):
        compact_text = re.sub(r'\s+', '', text)
        digits_only = re.sub(r'\D', '', text)

        if user_choices['id_num']:
            # CHỈ CHE NGUYÊN DÒNG NẾU DÒNG ĐÓ HOÀN TOÀN LÀ SỐ VÀ KHÔNG CÓ CHỮ CÁI
            if re.fullmatch(r'[\d\s.-]+', text.strip()) and 8 <= len(digits_only) <= 16:
                return True

        if user_choices['dob']:
            norm_text = normalize_text(text)
            has_dob_label = any(k in norm_text for k in ['ngay sinh', 'sinh ngay', 'date of birth', 'dob'])
            if has_dob_label and (
                re.search(r'\d{2}[\/\-.]\d{2}[\/\-.]\d{4}', text)
                or re.search(r'\d{4}[\/\-.]\d{2}[\/\-.]\d{2}', text)
            ):
                return True

        return False

    def redact_cv_portrait_area(img_ref):
        img_h_ref, img_w_ref = img_ref.shape[:2]
        pixelate_region(
            img_ref,
            int(img_w_ref * 0.07), int(img_h_ref * 0.10),
            int(img_w_ref * 0.36), int(img_h_ref * 0.33),
            blocks=12,
        )

    def merge_boxes(boxes, distance_x=8, distance_y=8):
        if not boxes:
            return []

        boxes = sorted(boxes, key=lambda b: (b[1], b[0]))
        merged = []

        for box in boxes:
            x1, y1, x2, y2 = box
            added = False

            for idx, (mx1, my1, mx2, my2) in enumerate(merged):
                overlap_y = y1 <= my2 + distance_y and y2 >= my1 - distance_y
                overlap_x = x1 <= mx2 + distance_x and x2 >= mx1 - distance_x

                if overlap_y and overlap_x:
                    merged[idx] = (min(mx1, x1), min(my1, y1), max(mx2, x2), max(my2, y2))
                    added = True
                    break

            if not added:
                merged.append(box)

        return merged

    def redact_cv_contact_smart(img_ref, ocr_items, line_items):
        img_h_ref, img_w_ref = img_ref.shape[:2]
        left_col_limit = int(img_w_ref * 0.45)
        email_pattern_cv = r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z0-9]{2,}(?:\.[a-zA-Z0-9]{2,})*'
        phone_pattern_cv = r'''(?<!\d)(?:(?:\+?84|0)[\s.-]?(?:\d[\s.-]?){8,10}|\d{3}[\s.-]?\d{3}[\s.-]?\d{4})(?!\d)'''

        def is_left_cv_item(item):
            return item['x1'] < left_col_limit

        def redact_email_rect(x1, y1, x2, y2):
            redact_rect(img_ref, max(0, x1 - 3), max(0, y1 - 7), min(img_w_ref, x2 + 3), min(img_h_ref, y2 + 4), padding=0)

        def redact_match_in_items(items, pattern, is_phone=False):
            items = [i for i in sorted(items, key=lambda k: k['x1']) if is_left_cv_item(i)]
            if not items:
                return False

            compact_text = ""
            parts = []
            for item in items:
                word = item['text'].strip()
                if not word:
                    continue

                word_compact = re.sub(r'\s+', '', word)
                start = len(compact_text)
                compact_text += word_compact
                end = len(compact_text)
                parts.append({'item': item, 'start': start, 'end': end})

            found = False
            for match in re.finditer(pattern, compact_text, re.VERBOSE if is_phone else 0):
                m_start, m_end = match.span()
                matched_items = [p['item'] for p in parts if p['end'] > m_start and p['start'] < m_end]
                if not matched_items:
                    continue

                x1 = min(i['x1'] for i in matched_items)
                y1 = min(i['y1'] for i in matched_items)
                x2 = max(i['x2'] for i in matched_items)
                y2 = max(i['y2'] for i in matched_items)

                if is_phone:
                    redact_rect(img_ref, x1, y1, x2, y2, padding=3)
                else:
                    redact_email_rect(x1, y1, x2, y2)

                found = True

            return found

        for item in ocr_items:
            if not is_left_cv_item(item):
                continue

            text = item['text'].strip()
            if not text:
                continue

            if re.search(email_pattern_cv, re.sub(r'\s+', '', text)):
                redact_email_rect(item['x1'], item['y1'], item['x2'], item['y2'])

            if re.search(phone_pattern_cv, text, re.VERBOSE):
                redact_rect(img_ref, item['x1'], item['y1'], item['x2'], item['y2'], padding=3)

        for line in line_items:
            redact_match_in_items(line['items'], email_pattern_cv)
            redact_match_in_items(line['items'], phone_pattern_cv, is_phone=True)

        email_label_lines = [
            line for line in line_items
            if re.search(r'\b(e\s*-?\s*mail|email|mail)\b', line['norm_text']) and line['x1'] < left_col_limit
        ]

        next_section_keywords = [
            'so thich', 'hobbies', 'ky nang', 'skills', 'hoc van',
            'education', 'kinh nghiem', 'experience', 'du an', 'projects'
        ]

        for email_line in email_label_lines:
            y_start = email_line['y2'] + 1
            y_end = min(img_h_ref, y_start + int(img_h_ref * 0.16))

            for line in line_items:
                if line['y1'] <= email_line['y2'] or line['x1'] >= left_col_limit:
                    continue
                if any(k in line['norm_text'] for k in next_section_keywords):
                    y_end = min(y_end, max(y_start + 1, line['y1'] - 3))
                    break

            items_below = [
                item for item in ocr_items
                if is_left_cv_item(item) and item['y1'] >= y_start and item['y2'] <= y_end
            ]
            if redact_match_in_items(items_below, email_pattern_cv):
                continue

            x_start = max(0, email_line['x1'] - 2)
            x_end = left_col_limit
            if y_end <= y_start or x_end <= x_start:
                continue

            roi = img_ref[y_start:y_end, x_start:x_end]
            if roi.size == 0:
                continue

            gray = cv2.cvtColor(roi, cv2.COLOR_BGR2GRAY)
            hsv = cv2.cvtColor(roi, cv2.COLOR_BGR2HSV)
            dark_mask = gray < 185
            saturated_mask = hsv[:, :, 1] > 35
            mask = (dark_mask | saturated_mask).astype(np.uint8) * 255
            mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN, cv2.getStructuringElement(cv2.MORPH_RECT, (3, 2)), iterations=1)
            mask = cv2.dilate(mask, cv2.getStructuringElement(cv2.MORPH_RECT, (9, 3)), iterations=1)

            contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            boxes = []
            for cnt in contours:
                x, y, w_box, h_box = cv2.boundingRect(cnt)
                if w_box < 12 or h_box < 3:
                    continue
                if h_box > (y_end - y_start) * 0.8:
                    continue
                boxes.append((x_start + x, y_start + y, x_start + x + w_box, y_start + y + h_box))

            for x1, y1, x2, y2 in merge_boxes(boxes, distance_x=10, distance_y=6):
                redact_email_rect(x1, y1, x2, y2)

    def redact_license_plates(img_ref, ocr_items):
        plate_patterns = [
            r'\d{2}[A-Z]{1,2}\d?[-.\s]?\d{3,5}[.]?\d{0,2}',
            r'\d{2}[-.\s]?[A-Z]{1,2}\d?[-.\s]?\d{3,5}[.]?\d{0,2}',
        ]

        for item in ocr_items:
            text = item['text'].upper()
            compact = re.sub(r'[^A-Z0-9]', '', text)
            looks_like_plate = (
                re.search(r'\d{2}[A-Z]{1,2}\d?', compact)
                and len(re.findall(r'\d', compact)) >= 5
                and 6 <= len(compact) <= 12
            )

            if looks_like_plate or any(re.search(p, text) for p in plate_patterns):
                redact_rect(img_ref, item['x1'] - 8, item['y1'] - 8, item['x2'] + 8, item['y2'] + 8, padding=0)
                item['is_plate'] = True
            else:
                item['is_plate'] = False

        remaining_items = sorted([i for i in ocr_items if not i.get('is_plate')], key=lambda i: i['y1'])

        for i in range(len(remaining_items) - 1):
            item1 = remaining_items[i]
            item2 = remaining_items[i + 1]

            if item2['y1'] - item1['y2'] > int(img_ref.shape[0] * 0.05):
                continue

            text = (item1['text'] + item2['text']).upper()
            compact = re.sub(r'[^A-Z0-9]', '', text)
            looks_like_plate = (
                re.search(r'\d{2}[A-Z]{1,2}\d?', compact)
                and len(re.findall(r'\d', compact)) >= 5
                and 6 <= len(compact) <= 11
            )

            if looks_like_plate:
                redact_rect(
                    img_ref,
                    min(item1['x1'], item2['x1']) - 8,
                    min(item1['y1'], item2['y1']) - 8,
                    max(item1['x2'], item2['x2']) + 8,
                    max(item1['y2'], item2['y2']) + 8,
                    padding=0,
                )

    # ==========================================
    # 6. OCR ITEM + GOM DÒNG
    # ==========================================
    ocr_items = []

    for bbox, easy_text, conf in pre_results:
        xs, ys = [p[0] for p in bbox], [p[1] for p in bbox]
        x1, y1 = int(min(xs) / scale_factor), int(min(ys) / scale_factor)
        x2, y2 = int(max(xs) / scale_factor), int(max(ys) / scale_factor)

        easy_text_clean = easy_text.strip()
        if not easy_text_clean:
            continue

        if '<<' in easy_text_clean or easy_text_clean.isdigit() or (conf > 0.85 and easy_text_clean.isupper()):
            final_text = easy_text_clean
        else:
            px1, py1 = max(0, x1 - 2), max(0, y1 - 2)
            px2, py2 = min(img_w, x2 + 2), min(img_h, y2 + 2)
            crop_img = img[py1:py2, px1:px2]

            if crop_img.shape[0] > 5 and crop_img.shape[1] > 5:
                pil_crop = Image.fromarray(cv2.cvtColor(crop_img, cv2.COLOR_BGR2RGB))
                vietocr_text = vietocr_reader.predict(pil_crop).strip()
                words = vietocr_text.split()
                final_text = easy_text_clean if (len(words) > 3 and len(set(words)) < len(words) / 2) else (vietocr_text or easy_text_clean)
            else:
                final_text = easy_text_clean

        ocr_items.append({
            'text': final_text,
            'norm_text': normalize_text(final_text),
            'conf': conf,
            'x1': x1,
            'y1': y1,
            'x2': x2,
            'y2': y2,
            'center_y': (y1 + y2) // 2,
            'height': y2 - y1,
        })

    ocr_items = sorted(ocr_items, key=lambda item: (item['y1'], item['x1']))

    print("\n===== OCR ITEMS SAU KHI XỬ LÝ =====")
    for idx, item in enumerate(ocr_items, start=1):
        print(f"{idx}. Text: {item['text']} | Norm: {item['norm_text']} | Box: ({item['x1']}, {item['y1']}, {item['x2']}, {item['y2']})")
    print("===== HẾT OCR ITEMS =====\n")

    lines = []
    for item in ocr_items:
        added = False
        for line in lines:
            if abs(item['center_y'] - line['center_y']) < max(18, max(12, item['height']) * 0.7):
                line['items'].append(item)
                line['center_y'] = sum(i['center_y'] for i in line['items']) // len(line['items'])
                added = True
                break

        if not added:
            lines.append({'center_y': item['center_y'], 'items': [item]})

    line_items = []
    for line in lines:
        items = sorted(line['items'], key=lambda item: item['x1'])
        line_text = ' '.join(item['text'] for item in items)
        line_items.append({
            'text': line_text,
            'norm_text': normalize_text(line_text),
            'x1': min(i['x1'] for i in items),
            'y1': min(i['y1'] for i in items),
            'x2': max(i['x2'] for i in items),
            'y2': max(i['y2'] for i in items),
            'center_y': line['center_y'],
            'height': max(i['y2'] for i in items) - min(i['y1'] for i in items),
            'items': items,
        })

    line_items = sorted(line_items, key=lambda item: item['y1'])

    print("\n===== LINE ITEMS ĐÃ GOM DÒNG =====")
    for idx, line in enumerate(line_items, start=1):
        print(f"{idx}. Line: {line['text']} | Norm: {line['norm_text']} | Box: ({line['x1']}, {line['y1']}, {line['x2']}, {line['y2']})")
    print("===== HẾT LINE ITEMS =====\n")

    # ==========================================
    # 7. TẠO REDACT BOXES CHO TEXT
    # ==========================================
    redact_boxes = []

    date_pattern = r'\d{1,2}[\/\-.\s]+\d{1,2}[\/\-.\s]+\d{2,4}'
    dob_keywords = ['ngay sinh', 'sinh ngay', 'date of birth', 'dob', 'nam sinh', 'ngay/thang/nam sinh']
    address_keywords = [
        'que quan', 'place of origin', 'noi sinh', 'place of birth',
        'noi thuong tru', 'thuong tru', 'place of residence', 'residence',
        'dia chi', 'address'
    ]
    stop_labels = [
        'ngay sinh', 'sinh ngay', 'date of birth', 'dob', 'nam sinh',
        'nganh', 'khoa hoc', 'khoa',
        'que quan', 'place of origin',
        'noi thuong tru', 'thuong tru', 'place of residence',
        'dia chi', 'address',
        'gioi tinh', 'sex', 'quoc tich', 'nationality',
        'ho ten', 'ho va ten', 'name',
        'dac diem nhan dang', 'personal identification', 'nhan dang'
    ]

    ADDRESS_LEFT_PAD = 8
    ADDRESS_RIGHT_PAD = 40
    ADDRESS_TOP_PAD = 2
    ADDRESS_BOTTOM_PAD = 6

    def next_line_is_another_label(line):
        return any(label in line['norm_text'] for label in stop_labels)

    def redact_driver_license_number():
        if not user_choices.get('driver_license_num'):
            return

        for index, line in enumerate(line_items):
            if 'so giay phep' not in line['norm_text']:
                continue

            for item in line['items']:
                digits_only = re.sub(r'\D', '', item['text'])
                if len(digits_only) >= 6:
                    redact_boxes.append((item['x1'], item['y1'], item['x2'], item['y2']))

            if index + 1 < len(line_items):
                next_line = line_items[index + 1]
                digits_only = re.sub(r'\D', '', next_line['text'])
                if len(digits_only) >= 6:
                    box_width = next_line['x2'] - next_line['x1']
                    redact_boxes.append((
                        next_line['x1'],
                        next_line['y1'],
                        next_line['x1'] + int(box_width * 0.5),
                        next_line['y2'],
                    ))

    def redact_cccd_address_by_items():
        if not user_choices['address']:
            return

        address_labels_found = []

        for item in ocr_items:
            item_norm = item['norm_text']
            if any(k in item_norm for k in ['que quan', 'place of origin']):
                address_labels_found.append({'type': 'origin', 'item': item})
            if any(k in item_norm for k in ['noi thuong tru', 'thuong tru', 'place of residence']):
                address_labels_found.append({'type': 'residence', 'item': item})

        if not address_labels_found:
            return

        address_labels_found = sorted(address_labels_found, key=lambda x: x['item']['y1'])

        for idx, label_data in enumerate(address_labels_found):
            label_item = label_data['item']
            start_y = label_item['y1'] - 5
            end_y = address_labels_found[idx + 1]['item']['y1'] + 10 if idx + 1 < len(address_labels_found) else img_h

            label_regex = r'(que quan[\s/]*place of origin|noi thuong tru[\s/]*place of residence|thuong tru[\s/]*residence|que quan|noi thuong tru|place of origin|place of residence|dia chi|address)\s*[:;/-]*'
            matches = list(re.finditer(label_regex, label_item['norm_text']))

            if matches:
                char_w = (label_item['x2'] - label_item['x1']) / max(1, len(label_item['norm_text']))

                if matches[0].start() >= 3:
                    redact_boxes.append((
                        max(0, label_item['x1'] - ADDRESS_LEFT_PAD),
                        max(0, label_item['y1'] - ADDRESS_TOP_PAD),
                        min(img_w, int(label_item['x1'] + matches[0].start() * char_w) - 5),
                        min(img_h, label_item['y2'] + ADDRESS_BOTTOM_PAD),
                    ))

                if len(label_item['norm_text']) - matches[-1].end() >= 2:
                    redact_boxes.append((
                        max(0, int(label_item['x1'] + matches[-1].end() * char_w) - 25),
                        max(0, label_item['y1'] - ADDRESS_TOP_PAD),
                        min(img_w, label_item['x2'] + ADDRESS_RIGHT_PAD),
                        min(img_h, label_item['y2'] + ADDRESS_BOTTOM_PAD),
                    ))

            for item in ocr_items:
                if item is label_item:
                    continue

                item_norm = item['norm_text']
                item_center_y = (item['y1'] + item['y2']) // 2

                if item_center_y < start_y or item_center_y > end_y:
                    continue

                if abs(item_center_y - ((label_item['y1'] + label_item['y2']) // 2)) < 20:
                    if item['x2'] <= label_item['x2'] + 5:
                        continue

                if any(k in item_norm for k in [
                    'que quan', 'place of origin', 'origin',
                    'noi thuong tru', 'thuong tru', 'place of residence', 'residence',
                    'gioi tinh', 'sex', 'quoc tich', 'nationality',
                    'date of expiry', 'co gia tri den', 'tri den',
                    'ngay sinh', 'date of birth', 'ho va ten', 'ho ten', 'full name'
                ]):
                    continue

                if item['x1'] < int(img_w * 0.25):
                    continue

                if re.search(r'[a-zA-Z0-9]', item_norm):
                    redact_boxes.append((
                        max(0, item['x1'] - ADDRESS_LEFT_PAD),
                        max(0, item['y1'] - ADDRESS_TOP_PAD),
                        min(img_w, item['x2'] + ADDRESS_RIGHT_PAD),
                        min(img_h, item['y2'] + ADDRESS_BOTTOM_PAD),
                    ))

    def redact_cccd_name_by_items():
        if not user_choices['name']:
            return

        y_name, y_dob = None, None

        for item in ocr_items:
            norm = item['norm_text']
            if y_name is None and any(k in norm for k in ['ho va ten', 'ho ten', 'full name']):
                y_name = item['y1']
            if y_dob is None and any(k in norm for k in ['ngay sinh', 'date of birth', 'dob']):
                y_dob = item['y1']

        if y_name is None:
            return

        limit_y = y_dob if (y_dob and y_dob > y_name) else y_name + int(img_h * 0.15)
        label_regex = r'(ho va ten.*?full name|ho ten.*?full name|ho va ten|ho ten|full name|name)'
        y_pad = 2

        for item in ocr_items:
            norm = item['norm_text']
            center_y = (item['y1'] + item['y2']) // 2

            if not (y_name - 10 <= center_y <= limit_y + 10):
                continue
            if item['x1'] < img_w * 0.25:
                continue

            text_alphas = ''.join(c for c in item['text'] if c.isalpha())
            has_label = re.search(label_regex, norm)

            if not has_label:
                if len(text_alphas) >= 2 and text_alphas.isupper() and not any(k in norm for k in ['ngay', 'sinh', 'date', 'birth']):
                    redact_boxes.append((item['x1'] - 2, item['y1'] + y_pad, item['x2'] + 3, item['y2'] - y_pad))
                continue

            match = re.search(label_regex + r'\s*[:;/\-\|]*', norm)
            if match and len(norm) - match.end() >= 2:
                char_w = (item['x2'] - item['x1']) / max(1, len(norm))
                val_x1 = int(item['x1'] + match.end() * char_w) + 8
                if val_x1 < item['x2']:
                    redact_boxes.append((val_x1, item['y1'] + y_pad, item['x2'] + 3, item['y2'] - y_pad))

    def add_dob_boxes(index, line):
        text = line['text']
        matches = list(re.finditer(date_pattern, text))

        if matches:
            for match in matches:
                start_char, end_char = match.span()
                char_width = (line['x2'] - line['x1']) / max(1, len(text))
                dob_x1 = int(line['x1'] + start_char * char_width) - 130
                dob_x2 = int(line['x1'] + end_char * char_width) + 35
                line_height = max(12, line['height'])
                dob_y1 = line['y1'] - 2
                dob_y2 = min(line['y2'], line['y1'] + int(line_height * 0.85)) + 8
                redact_rect(img, max(0, dob_x1), max(0, dob_y1), min(img_w, dob_x2), min(img_h, dob_y2), padding=0)
            return

        for j in range(index + 1, min(index + 2, len(line_items))):
            next_line = line_items[j]
            if any(label in next_line['norm_text'] for label in stop_labels):
                break
            if re.search(date_pattern, next_line['text']):
                redact_boxes.append((
                    max(0, next_line['x1'] - 130),
                    max(0, next_line['y1'] - 2),
                    min(img_w, next_line['x2'] + 35),
                    min(img_h, next_line['y2'] + 2),
                ))

    # Xử lý riêng các phần dễ bị che nhầm.
    redact_cccd_address_by_items()
    redact_cccd_name_by_items()
    redact_driver_license_number()

    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(0|84|\+84)[3|5|7|8|9][0-9]{8}\b'

    for index, line in enumerate(line_items):
        is_sensitive = False
        norm = line['norm_text']
        text = line['text']

        # Không che các dòng ngày xác thực / hạn dùng của thẻ BHYT
        if document_type == "Thẻ Bảo hiểm y tế":
            if any(k in norm for k in [
                'ngay sinh',
                'sinh ngay',
                'ngay xac thuc',
                'xac thuc',
                'ngay ky',
                'ngay cap',
                'tu ngay',
                'den ngay',
                'gia tri su dung',
                'han su dung',
                'valid',
                'valid from',
                'valid until',
            ]):
                continue

        if user_choices['id_num']:
            for item in line['items']:
                item_text = item['text']
                compact_item = re.sub(r'\s+', '', item_text)
                
                # Tìm chuỗi có từ 9 đến 13 số liên tiếp (Mã SV, CMND, CCCD)
                if re.search(r'\d{9,13}', compact_item):
                    alphas = sum(c.isalpha() for c in item_text)
                    
                    # TRƯỜNG HỢP 1: Item này tách biệt (Chỉ có số, ít chữ cái) -> Che cả hộp item
                    if len(item_text) <= 16 or alphas < 3:
                        redact_boxes.append((item['x1'], item['y1'], item['x2'], item['y2']))
                        
                    # TRƯỜNG HỢP 2: AI (OCR) đọc dính liền cả họ tên và số -> Tính toán tọa độ riêng cho dải số
                    else:
                        match = re.search(r'\d[\d\s.-]{7,14}\d', item_text)
                        if match:
                            start_c, end_c = match.span()
                            char_w = (item['x2'] - item['x1']) / max(1, len(item_text))
                            redact_boxes.append((
                                int(item['x1'] + start_c * char_w) - 5,
                                item['y1'],
                                int(item['x1'] + end_c * char_w) + 5,
                                item['y2']
                            ))

        if user_choices.get('address') and any(k in norm for k in address_keywords):
            continue

        if not user_choices.get('address') and any(k in norm for k in address_keywords):
            continue

        if any(kw in norm for kw in ['het han', 'nganh', 'khoa hoc', 'nien khoa', 'truong dai hoc', 'the sinh vien']):
            continue

        if not user_choices['dob'] and (any(kw in norm for kw in dob_keywords) or re.search(date_pattern, text)):
            continue

        if '<<' in text:
            has_digits = any(char.isdigit() for char in text)
            if user_choices['name'] and not has_digits:
                redact_boxes.append((line['x1'], line['y1'], line['x2'], line['y2']))
            elif (user_choices['id_num'] or user_choices['dob']) and has_digits:
                redact_boxes.append((line['x1'], line['y1'], line['x2'], line['y2']))
            continue

        safe_headers = [
            'cong hoa xa hoi', 'chu nghia viet nam', 'doc lap tu do', 'hanh phuc',
            'can cuoc cong dan', 'citizen identity card', 'giam doc cong an'
        ]
        if any(safe in norm for safe in safe_headers) and not line_has_sensitive_pattern(text):
            continue

        if user_choices['id_num'] and any(k in norm for k in ['so i no', 'so / no', 'số / no', 'no']):
            for item in line['items']:
                digits_only = re.sub(r'\D', '', item['text'])
                if len(digits_only) >= 9:
                    redact_boxes.append((item['x1'], item['y1'], item['x2'], item['y2']))
            continue

        if user_choices['dob'] and any(kw in norm for kw in dob_keywords):
            add_dob_boxes(index, line)
            continue

        if user_choices['cv_contact'] and (re.search(email_pattern, text) or re.search(phone_pattern, text)):
            is_sensitive = True

        

        is_id_label_line = user_choices['id_num'] and any(k in norm for k in ['so i no', 'so / no', 'số / no'])
        is_address_related_line = user_choices['address'] and any(k in norm for k in address_keywords)

        if line_has_sensitive_pattern(text) or (has_keyword(norm, sensitive_keywords) and not is_id_label_line and not is_address_related_line):
            is_sensitive = True

        if has_keyword(norm, next_line_keywords):
            for j in range(index + 1, min(index + 2, len(line_items))):
                next_norm = line_items[j]['norm_text']
                if not user_choices.get('address') and any(k in next_norm for k in address_keywords):
                    continue
                if not next_line_is_another_label(line_items[j]):
                    redact_boxes.append((line_items[j]['x1'], line_items[j]['y1'], line_items[j]['x2'], line_items[j]['y2']))

        if is_sensitive:
            redact_boxes.append((line['x1'], line['y1'], line['x2'], line['y2']))

    print("\n===== REDACT BOXES SẼ CHE =====")
    for idx, box in enumerate(redact_boxes, start=1):
        print(f"{idx}. Box: {box}")
    print("===== HẾT REDACT BOXES =====\n")

    for x1, y1, x2, y2 in redact_boxes:
        redact_rect(img, max(0, x1 - 5), max(0, y1 - 3), min(img_w, x2 + 5), min(img_h, y2 + 3), padding=0)

    # ==========================================
    # 8. LỚP MỜ HÌNH ẢNH
    # ==========================================
    if user_choices['qr']:
        qr_y_limit = int(img_h * 0.4)
        for item in ocr_items:
            if any(k in item['norm_text'] for k in ['can cuoc', 'citizen', 'so /', 'no.']):
                qr_y_limit = max(qr_y_limit, item['y2'])

        pixelate_region(
            img,
            int(img_w * 0.8),
            int(img_w * 0.05),
            img_w - 20,
            qr_y_limit - 60,
            blocks=15,
        )

    if user_choices['finger']:
        anchor_groups = {
            'left': ['ngon tro trai', 'left index', 'tro trai'],
            'right': ['ngon tro phai', 'right index', 'tro phai'],
        }

        finger_boxes = []
        for item in ocr_items:
            norm = item['norm_text']
            x1, y1, x2, y2 = item['x1'], item['y1'], item['x2'], item['y2']
            h_text = y2 - y1

            if any(anchor in norm for anchors in anchor_groups.values() for anchor in anchors):
                finger_boxes.append({
                    'x1': max(0, x1 - int(h_text * 1.5)),
                    'y1': max(0, y1 - int(h_text * 9.0)),
                    'x2': min(img_w, x2 + int(h_text * 1.5)),
                    'y2': max(0, y1 - 2),
                })

        for box in finger_boxes:
            pixelate_region(img, box['x1'], box['y1'], box['x2'], box['y2'], blocks=8)

    if user_choices['face']:
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray_for_face = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        faces = face_cascade.detectMultiScale(
            gray_for_face,
            scaleFactor=1.1,
            minNeighbors=5,
            minSize=(40, 40),
        )

        for (x, y, w_face, h_face) in faces:
            pixelate_region(img, x, y, x + w_face, y + h_face, blocks=8)

        if document_type.startswith("CV"):
            redact_cv_portrait_area(img)
        elif document_type in ["Căn cước công dân (Mặt trước)", "Thẻ sinh viên / Thẻ học sinh"]:
            pixelate_region(img, int(img_w * 0.03), int(img_h * 0.4), int(img_w * 0.30), int(img_h * 0.82), blocks=10)

    if user_choices['cv_contact']:
        redact_cv_contact_smart(img, ocr_items, line_items)

    if user_choices.get('plate'):
        redact_license_plates(img, ocr_items)

    if user_choices['barcode']:
        try:
            pixelate_region(
                img,
                int(img_w * 0.35),
                int(img_h * 0.83),
                int(img_w * 0.92),
                int(img_h * 0.98),
                blocks=10,
            )
        except Exception as e:
            print(f"Lỗi khi che mã vạch (đã bỏ qua): {e}")

    # --- 9. LƯU ẢNH HỖ TRỢ ĐƯỜNG DẪN TIẾNG VIỆT ---
    try:
        out_ext = os.path.splitext(output_path)[1].lower()
        
        # NẾU ĐẦU RA YÊU CẦU LÀ PDF -> DÙNG THƯ VIỆN PIL ĐỂ LƯU
        if out_ext == '.pdf':
            # Chuyển hệ màu OpenCV (BGR) sang PIL (RGB)
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pil_pdf = Image.fromarray(img_rgb)
            # Lưu ảnh thành 1 trang PDF
            pil_pdf.save(output_path, "PDF", resolution=100.0)
            print(f"✨ Đã lưu file PDF an toàn vào: {output_path}")
            return True

        # NẾU LÀ ẢNH BÌNH THƯỜNG -> DÙNG OPENCV ĐỂ LƯU
        if out_ext not in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
            out_ext = '.jpg' # Mặc định an toàn
            
        img = np.ascontiguousarray(img, dtype=np.uint8)
        is_success, buffer = cv2.imencode(out_ext, img)

        if not is_success:
            print(f"Lỗi hệ thống: Không thể mã hóa sang định dạng {out_ext}.")
            return False

        with open(output_path, "wb") as f:
            f.write(buffer.tobytes())

        print(f"✨ Đã lưu file ảnh thành công vào: {output_path}")
        return True

    except Exception as e:
        print(f"Lỗi ghi file: {e}")
        return False
    
# ==========================================
# TRUNG TÂM PHÂN LOẠI FILE (ROUTER)
# ==========================================

# HÀM XỬ LÝ RIÊNG CHO FILE WORD (GIỮ NGUYÊN ĐỊNH DẠNG TEXT)
# ==========================================
def process_and_redact_docx(file_path, output_path, parent_window):
    # 1. HỘP THOẠI CHỌN VÙNG CẦN CHE (Rút gọn cho Text)
    user_choices = {'id_num': False, 'cv_contact': False, 'dob': False}
    is_submitted = [False]

    dialog = tk.Toplevel(parent_window)
    dialog.title("Che thông tin file Word")
    dialog.geometry("350x250")
    dialog.configure(bg="#F0F7F4")
    dialog.transient(parent_window)
    dialog.grab_set()

    tk.Label(dialog, text="🛠️ ĐIỀU CHỈNH VÙNG CẦN CHE:", bg="#F0F7F4", font=("Arial", 11, "bold")).pack(pady=10)

    vars_dict = {}
    options = [
        ('id_num', "Che Số định danh / Mã sinh viên"),
        ('cv_contact', "Che Email / Số điện thoại"),
        ('dob', "Che Ngày sinh")
    ]

    for key, label in options:
        var = tk.BooleanVar(value=True) # Mặc định tick sẵn
        vars_dict[key] = var
        tk.Checkbutton(dialog, text=label, variable=var, bg="#F0F7F4").pack(anchor="w", padx=40, pady=5)

    def on_submit():
        for key in user_choices.keys():
            user_choices[key] = vars_dict[key].get()
        is_submitted[0] = True
        dialog.destroy()

    tk.Button(dialog, text="Xác nhận & Xử lý Word", command=on_submit, bg="#A8DADC", font=("Arial", 10, "bold"), padx=15).pack(pady=20)
    
    parent_window.wait_window(dialog)

    if not is_submitted[0]:
        return False

    # 2. ĐỌC VÀ THAY THẾ TEXT TRONG FILE WORD
    try:
        doc = docx.Document(file_path)
        
        def redact_text(text):
            # Thay thế Mã số (9-13 số)
            if user_choices['id_num']:
                text = re.sub(r'\b\d{9,13}\b', '***', text)
            # Thay thế SĐT và Email
            if user_choices['cv_contact']:
                text = re.sub(r'(0|84|\+84)[3|5|7|8|9][0-9]{8}\b', '[ĐÃ CHE SĐT]', text)
                text = re.sub(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z0-9]{2,}', '[ĐÃ CHE EMAIL]', text)
            # Thay thế Ngày sinh
            if user_choices['dob']:
                text = re.sub(r'\b\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4}\b', '[ĐÃ CHE NGÀY SINH]', text)
            return text
        # HÀM MỚI: Duyệt qua từng run để không làm mất ảnh
        def process_paragraphs(paragraphs):
            for para in paragraphs:
                for run in para.runs:
                    # Kiểm tra xem run có chứa ảnh không (thường là InlineShapes)
                    # Nếu run không phải là ảnh, mới thay thế text
                    if not run.element.xpath('.//a:blip'): # Kiểm tra XML của ảnh
                        run.text = redact_text(run.text)
        # Quét các đoạn văn bản thường
        process_paragraphs(doc.paragraphs)
            
        # Quét các đoạn văn bản nằm trong Bảng (Table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = redact_text(para.text)

        doc.save(output_path)
        print(f"✨ Đã lưu file Word thành công vào: {output_path}")
        return True

    except Exception as e:
        messagebox.showerror("Lỗi Word", f"Không thể xử lý file Word:\n{e}", parent=parent_window)
        return False
def handle_input_file(file_path, output_path, parent_window):
    """Xử lý ảnh (jpg/png) và file Word (doc/docx)"""
    ext = file_path.lower().split('.')[-1]
    
    # 1. TRƯỜNG HỢP GỬI FILE WORD
    if ext in ['doc', 'docx']:
        print("🔄 Xử lý file Word...")
        out_docx_path = output_path.rsplit('.', 1)[0] + ".docx"
        is_success = process_and_redact_docx(file_path, out_docx_path, parent_window)
        if is_success:
            # Nếu file đã đổi tên, ta trả về đường dẫn mới cho SocialHub
            if os.path.exists(out_docx_path):
                return True
        return False

    # 2. TRƯỜNG HỢP LÀ ẢNH GỐC
    elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'webp']:
        return process_and_redact(file_path, output_path, parent_window)
        
    else:
        messagebox.showerror("Định dạng", f"Định dạng .{ext} không được hỗ trợ!", parent=parent_window)
        return False